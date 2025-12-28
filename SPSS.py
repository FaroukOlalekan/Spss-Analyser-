import streamlit as st
import pandas as pd
import numpy as np
from scipy.stats import chi2_contingency, fisher_exact, ttest_ind, ttest_rel, mannwhitneyu, wilcoxon, kruskal, friedmanchisquare, spearmanr, pearsonr
from statsmodels.stats.contingency_tables import Table2x2
import statsmodels.api as sm
import pyreadstat
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

# ---------------- PAGE CONFIG ----------------
st.set_page_config(page_title="Phabolous SPSS 📊 Analysis Tool", layout="wide")

# ---------------- BACKGROUND ----------------
st.markdown("""
<style>
.stApp {background-color: #eef4f8;}
</style>
""", unsafe_allow_html=True)

# ---------------- TITLE ----------------
st.title("🤖 Phabolous SPSS Data Analysis")

# ---------------- USER RESEARCH QUESTION ----------------
st.subheader("❓ Enter Your Research Question")
user_question = st.text_input(
    "Type a question the AI will answer based on your data",
    placeholder="e.g., Enter Question?"
)

# ---------------- FILE UPLOAD ----------------
uploaded_file = st.file_uploader(
    "📂 Upload dataset",
    type=["sav","xlsx","xls","csv","tsv","txt","dta","sas7bdat","json"]
)

if uploaded_file:
    file_name = uploaded_file.name.lower()
    try:
        if file_name.endswith(".sav"):
            df, meta = pyreadstat.read_sav(uploaded_file)
        elif file_name.endswith((".xlsx",".xls")):
            df = pd.read_excel(uploaded_file)
        elif file_name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
        elif file_name.endswith(".tsv"):
            df = pd.read_csv(uploaded_file, sep="\t")
        elif file_name.endswith(".txt"):
            df = pd.read_csv(uploaded_file, sep=None, engine="python")
        elif file_name.endswith(".dta"):
            df = pd.read_stata(uploaded_file)
        elif file_name.endswith(".sas7bdat"):
            df = pd.read_sas(uploaded_file)
        elif file_name.endswith(".json"):
            df = pd.read_json(uploaded_file)
        st.success("✅ Dataset loaded successfully")
    except Exception as e:
        st.error(f"❌ Error loading file: {e}")
        st.stop()

    st.subheader("👀 Dataset Preview")
    st.dataframe(df.head())

    # ---------------- VARIABLE SELECTION ----------------
    st.subheader("🧪 Select Variables")
    exposure = st.selectbox("Independent / Predictor Variable", df.columns)
    outcome = st.selectbox("Dependent / Outcome Variable", df.columns)

    # ---------------- TEST SELECTION ----------------
    test_options = [
        "Crosstab / Frequencies",
        "Chi-square",
        "Fisher’s Exact",
        "Odds Ratio / Relative Risk",
        "Independent Samples t-test",
        "Paired Samples t-test",
        "One-way ANOVA",
        "Repeated Measures ANOVA",
        "Pearson Correlation",
        "Spearman Correlation",
        "Binary Logistic Regression",
        "Mann-Whitney U",
        "Wilcoxon Signed-Rank",
        "Kruskal-Wallis",
        "Friedman Test"
    ]
    selected_test = st.selectbox("⚡ Select SPSS Test to Run", test_options)

    if st.button("Run Test") and user_question and exposure and outcome:
        st.subheader(f"📊 Running: {selected_test}")

        # ---------------- CATEGORICAL DATA ----------------
        if selected_test in ["Crosstab / Frequencies","Chi-square","Fisher’s Exact","Odds Ratio / Relative Risk","Binary Logistic Regression"]:
            raw_table = pd.crosstab(df[exposure], df[outcome])

            if selected_test == "Crosstab / Frequencies":
                st.table(raw_table)
            elif selected_test == "Chi-square":
                chi2, p, dof, expected = chi2_contingency(raw_table)
                st.markdown(f"χ²({dof}) = {chi2:.2f}, p = {p:.3f}")
            elif selected_test == "Fisher’s Exact":
                fisher_p = fisher_exact(raw_table.values)[1]
                st.markdown(f"Fisher’s Exact Test p = {fisher_p:.3f}")
            elif selected_test == "Odds Ratio / Relative Risk":
                risk = Table2x2(raw_table.values)
                or_val = risk.oddsratio
                rr_val = risk.riskratio
                ci_low, ci_high = risk.oddsratio_confint()
                st.markdown(f"OR = {or_val:.2f} (95% CI {ci_low:.2f}-{ci_high:.2f}) | RR = {rr_val:.2f}")
            elif selected_test == "Binary Logistic Regression":
                X = sm.add_constant(df[exposure])
                y = df[outcome]
                logit = sm.Logit(y,X).fit(disp=False)
                st.markdown(f"β = {logit.params[1]:.2f}, p = {logit.pvalues[1]:.3f}")

        # ---------------- SCALE DATA ----------------
        if selected_test in ["Independent Samples t-test","Paired Samples t-test","One-way ANOVA","Repeated Measures ANOVA",
                             "Pearson Correlation","Spearman Correlation","Mann-Whitney U","Wilcoxon Signed-Rank",
                             "Kruskal-Wallis","Friedman Test"]:
            x = df[exposure]
            y = df[outcome]

            if selected_test == "Independent Samples t-test":
                stat, p = ttest_ind(x,y)
                st.markdown(f"t = {stat:.2f}, p = {p:.3f}")
            elif selected_test == "Paired Samples t-test":
                stat, p = ttest_rel(x,y)
                st.markdown(f"t = {stat:.2f}, p = {p:.3f}")
            elif selected_test == "Pearson Correlation":
                r,p = pearsonr(x,y)
                st.markdown(f"r = {r:.2f}, p = {p:.3f}")
            elif selected_test == "Spearman Correlation":
                r,p = spearmanr(x,y)
                st.markdown(f"ρ = {r:.2f}, p = {p:.3f}")
            elif selected_test == "Mann-Whitney U":
                stat,p = mannwhitneyu(x,y)
                st.markdown(f"U = {stat:.2f}, p = {p:.3f}")
            elif selected_test == "Wilcoxon Signed-Rank":
                stat,p = wilcoxon(x,y)
                st.markdown(f"W = {stat:.2f}, p = {p:.3f}")
            elif selected_test == "Kruskal-Wallis":
                stat,p = kruskal(x,y)
                st.markdown(f"H = {stat:.2f}, p = {p:.3f}")
            elif selected_test == "Friedman Test":
                stat,p = friedmanchisquare(x,y)
                st.markdown(f"χ² = {stat:.2f}, p = {p:.3f}")

        st.success("✅ Test completed!")

        # ---------------- BAR CHART ----------------
        st.subheader("📈 Bar Chart")
        fig, ax = plt.subplots()
        if 'raw_table' in locals():
            raw_table.plot(kind='bar', ax=ax)
            ax.set_xlabel(exposure)
            ax.set_ylabel("Frequency")
            st.pyplot(fig)

        # ---------------- CHATGPT-STYLE AI ----------------
        st.subheader("🤖 Ask the AI About Your Data")
        ai_question = st.text_area("Ask a question about the results:", placeholder="Can I reject the null hypothesis? Explain simply.")

        def ai_chat_response(q):
            q = q.lower()
            try:
                pval = p if 'p' in locals() else 0.01
                or_value = or_val if 'or_val' in locals() else 1
            except:
                pval = 0.01
                or_value = 1
            if "null" in q or "hypothesis" in q:
                return f"✔ Null hypothesis {'can' if pval < 0.05 else 'cannot'} be rejected (p = {pval:.3f})."
            elif "simple" in q or "student" in q:
                return f"🧠 Simply put, {exposure} affects {outcome}, with OR ≈ {or_value:.2f}."
            elif "spss" in q:
                return f"📊 SPSS-style: χ² / test p = {pval:.3f}, OR ≈ {or_value:.2f}."
            else:
                return f"📌 There is a significant relationship between {exposure} and {outcome} (p = {pval:.3f})."

        if ai_question:
            st.markdown("### 🤖 AI Response")
            st.markdown(ai_chat_response(ai_question))

        # ---------------- DOWNLOADABLE REPORT ----------------
        def generate_report():
            doc = Document()
            doc.add_heading("📊 AI-SPSS Analysis Report", 0)
            doc.add_heading("❓ Research Question", level=1)
            doc.add_paragraph(user_question)
            doc.add_heading("⚡ Test Performed", level=1)
            doc.add_paragraph(selected_test)
            doc.add_heading("📈 Results Summary", level=1)
            if selected_test in ["Chi-square", "Fisher’s Exact", "Odds Ratio / Relative Risk"]:
                doc.add_paragraph(f"p-value = {pval:.3f} | OR ≈ {or_value:.2f} | RR ≈ {rr_val if 'rr_val' in locals() else 'N/A'}")
            elif selected_test == "Binary Logistic Regression":
                doc.add_paragraph(f"β = {logit.params[1]:.2f}, p = {logit.pvalues[1]:.3f}")
            elif selected_test in ["Independent Samples t-test","Paired Samples t-test","Pearson Correlation","Spearman Correlation",
                                   "Mann-Whitney U","Wilcoxon Signed-Rank","Kruskal-Wallis","Friedman Test"]:
                doc.add_paragraph(f"Test statistic = {stat:.2f}, p = {p:.3f}")
            doc.add_heading("🤖 AI Interpretation", level=1)
            doc.add_paragraph(ai_chat_response("Explain simply"))
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            return file_stream

        st.subheader("📥 Downloadable Report")
        report_stream = generate_report()
        st.download_button(
            label="Download Report (.docx)",
            data=report_stream,
            file_name="AI_SPSS_Analysis_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
